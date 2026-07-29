#!/usr/bin/env python3
"""
PRE-UPLOAD AUDIT — Masters X compiled manuscripts
Checks: editorial fixes present/absent, italic preservation, cross-format consistency.
"""
from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document

try:
    import fitz
except ImportError:
    raise SystemExit("PyMuPDF (fitz) required")

ROOT = Path(__file__).resolve().parents[2]
STAGING = ROOT / "production_staging"
OUT = STAGING / "PRE_UPLOAD_AUDIT_REPORT.md"

# Required NEW strings by book (from FIX_CHANGELOG + apartment pass + dwelling follow-ups)
REQUIRED = {
    1: [
        "1647 Genessee",
        "Warren County",
        "Osage ancestral",
        "Midwest Precote",
        "Hunt Midwest",
        "RIVERWARDS",
        "hundred and fifty feet",
        "bottomed out at a hundred and sixty",
    ],
    2: [
        "Washington Street office",
        "Pennsylvania Avenue",
        "Washington Street entrance",
        "mailbox. Washington Street",
        "mailbox on Washington Street",
        "11 PM. Washington Street",
        "apartment hummed its limestone",
        "bathroom sink in an apartment in Kansas City",
    ],
    3: [
        "Iceland basalt chamber",
        "Across town, 160 feet below the Northland bluff",
        "streets of Quality Hill empty",
        "across the river at SubTropolis",
        "She drove the long way",
        "third floor of the Washington Street building",
        "Quality Hill balcony",
        "The apartment hummed at 55 Hz",
        "apartment was clean. The apartment was empty",
        "hallway outside his door",  # Marcus Chen follow-up
    ],
}

# Banned OLD strings (must be absent). Book-scoped where needed.
BANNED = {
    1: [
        "2847 Genessee",
        "Washington County, Missouri",
        "Bethany Falls Limestone Company",
        "MISSOURI COLD STORAGE",
        "fifty-five feet below the surface",
        "The shaft dropped a hundred and sixty feet",
    ],
    2: [
        "Hotel Phillips Building office",
        "buildings of Troost Avenue",
        "Quality Hill entrance on graduation",
        "mailbox. Quality Hill. February",
        "mailbox on Quality Hill held",
        "11 PM. Quality Hill. The Foundation",
        "The house hummed its limestone fundamental",
        "bathroom sink in a house in Kansas City",
    ],
    3: [
        "Somewhere below them, 160 feet below, in the SubTropolis",
        "Troost corridor empty, the industrial district",
        "West Bottoms, where, 160 feet below ground, the Moreau",
        "She drove down Troost. Past the gas station",
        "second floor of the Washington Street building",
        "The house was clean. The house was empty",
        "The house hummed at 55 Hz",
        "He sat on the Quality Hill porch",
        "Marcus Chen on his front porch",
        # residual that should be gone after apartment pass
        "The house hummed at 53.6",
    ],
}

# Formats to audit
FORMATS = {
    "BUILD_DOCX_1": STAGING / "_sources/build_docx/MASTERS_X_BOOK1_BUILD.docx",
    "BUILD_DOCX_2": STAGING / "_sources/build_docx/MASTERS_X_BOOK2_BUILD.docx",
    "BUILD_DOCX_3": STAGING / "_sources/build_docx/MASTERS_X_BOOK3_BUILD.docx",
    "SRC_DOCX_1": STAGING / "_sources/MASTERS_X_BOOK1_ITALICIZED_FIXED.docx",
    "SRC_DOCX_2": STAGING / "_sources/MASTERS_X_BOOK2_ITALICIZED_FIXED.docx",
    "SRC_DOCX_3": STAGING / "_sources/MASTERS_X_BOOK3_ITALICIZED_FIXED.docx",
    "OMNI_HC": STAGING / "omnibus/9798295884412_HC/interior.pdf",
    "OMNI_PB": STAGING / "omnibus/9798256072704_PB/interior.pdf",
    "B1_HC": STAGING / "b1_inheritance/9798295800801_HC/interior.pdf",
    "B1_PB": STAGING / "b1_inheritance/9798256008048_PB/interior.pdf",
    "B2_HC": STAGING / "b2_grimoire/9798295812675_HC/interior.pdf",
    "B2_PB": STAGING / "b2_grimoire/9798256009953_PB/interior.pdf",
    "B3_HC": STAGING / "b3_kingdom/9798295812705_HC/interior.pdf",
    "B3_PB": STAGING / "b3_kingdom/9798256010072_PB/interior.pdf",
    "B1_EPUB": STAGING / "b1_inheritance/9798256008819_EPUB/9798256008819.epub",
    "B2_EPUB": STAGING / "b2_grimoire/9798256009625_EPUB/9798256009625.epub",
    "B3_EPUB": STAGING / "b3_kingdom/9798256009809_EPUB/9798256009809.epub",
}

BOOK_OF = {
    "BUILD_DOCX_1": 1, "SRC_DOCX_1": 1, "B1_HC": 1, "B1_PB": 1, "B1_EPUB": 1,
    "BUILD_DOCX_2": 2, "SRC_DOCX_2": 2, "B2_HC": 2, "B2_PB": 2, "B2_EPUB": 2,
    "BUILD_DOCX_3": 3, "SRC_DOCX_3": 3, "B3_HC": 3, "B3_PB": 3, "B3_EPUB": 3,
    "OMNI_HC": 0, "OMNI_PB": 0,  # all books
}


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"'))


def load_docx_text(path: Path) -> str:
    doc = Document(str(path))
    return norm("\n".join(p.text or "" for p in doc.paragraphs))


def load_docx_italic_stats(path: Path) -> dict:
    doc = Document(str(path))
    italic_runs = italic_chars = total_runs = 0
    paras_with_italic = 0
    for p in doc.paragraphs:
        hit = False
        for r in p.runs:
            total_runs += 1
            t = r.text or ""
            if r.italic and t.strip():
                italic_runs += 1
                italic_chars += len(t)
                hit = True
        if hit:
            paras_with_italic += 1
    return {
        "italic_runs": italic_runs,
        "italic_chars": italic_chars,
        "paras_with_italic": paras_with_italic,
        "total_runs": total_runs,
    }


def load_pdf_text(path: Path) -> str:
    doc = fitz.open(path)
    parts = []
    for i in range(doc.page_count):
        parts.append(doc[i].get_text())
    doc.close()
    return norm("\n".join(parts))


def load_pdf_italic_stats(path: Path, sample_pages: int = 40) -> dict:
    """Count spans whose font name suggests italic/oblique."""
    doc = fitz.open(path)
    italic_spans = 0
    total_spans = 0
    italic_chars = 0
    step = max(1, doc.page_count // sample_pages)
    pages_sampled = 0
    for i in range(0, doc.page_count, step):
        pages_sampled += 1
        d = doc[i].get_text("dict")
        for b in d.get("blocks", []):
            for line in b.get("lines", []):
                for s in line.get("spans", []):
                    total_spans += 1
                    font = (s.get("font") or "").lower()
                    flags = s.get("flags", 0)
                    is_it = ("italic" in font) or ("oblique" in font) or bool(flags & 2)
                    if is_it and (s.get("text") or "").strip():
                        italic_spans += 1
                        italic_chars += len(s.get("text") or "")
    doc.close()
    return {
        "italic_spans": italic_spans,
        "italic_chars": italic_chars,
        "total_spans": total_spans,
        "pages_sampled": pages_sampled,
    }


def load_epub_text(path: Path) -> str:
    parts = []
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if n.endswith((".xhtml", ".html", ".htm")):
                raw = z.read(n).decode("utf-8", errors="ignore")
                # strip tags lightly
                raw = re.sub(r"<[^>]+>", " ", raw)
                parts.append(raw)
    return norm("\n".join(parts))


def load_epub_italic_stats(path: Path) -> dict:
    em_count = i_count = 0
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if n.endswith((".xhtml", ".html", ".htm")):
                raw = z.read(n).decode("utf-8", errors="ignore")
                em_count += len(re.findall(r"<em\b", raw, re.I))
                i_count += len(re.findall(r"<i\b", raw, re.I))
    return {"em_tags": em_count, "i_tags": i_count, "italic_tags": em_count + i_count}


def check_strings(text: str, required: list[str], banned: list[str]) -> dict:
    req = {s: (norm(s) in text) for s in required}
    ban = {s: (norm(s) in text) for s in banned}
    return {
        "required_pass": all(req.values()),
        "banned_clean": not any(ban.values()),
        "required": req,
        "banned_hits": {k: v for k, v in ban.items() if v},
        "required_misses": [k for k, v in req.items() if not v],
    }


def main() -> int:
    lines = []
    lines.append("# PRE-UPLOAD AUDIT REPORT")
    lines.append("**Date:** 2026-07-28")
    lines.append("**Scope:** Editorial fixes · italics · cross-format consistency")
    lines.append("**Authority:** `FIX_CHANGELOG.md` + apartment/balcony pass + dwelling follow-ups")
    lines.append("")

    texts: dict[str, str] = {}
    results: dict[str, dict] = {}

    # Load texts
    for key, path in FORMATS.items():
        if not path.exists():
            results[key] = {"error": f"MISSING {path}"}
            continue
        if path.suffix.lower() == ".docx":
            texts[key] = load_docx_text(path)
        elif path.suffix.lower() == ".pdf":
            texts[key] = load_pdf_text(path)
        elif path.suffix.lower() == ".epub":
            texts[key] = load_epub_text(path)

    # --- Section 1: Editorial fixes ---
    lines.append("## 1. Editorial corrections (required present / banned absent)")
    lines.append("")
    overall_editorial_ok = True
    for key, path in FORMATS.items():
        if key.startswith("SRC_DOCX"):
            continue  # stale sources expected to fail required
        if key not in texts:
            lines.append(f"- **{key}**: MISSING FILE")
            overall_editorial_ok = False
            continue
        book = BOOK_OF[key]
        if book == 0:
            req = REQUIRED[1] + REQUIRED[2] + REQUIRED[3]
            ban = BANNED[1] + BANNED[2] + BANNED[3]
        else:
            req = REQUIRED[book]
            ban = BANNED[book]
        r = check_strings(texts[key], req, ban)
        results[key] = r
        status = "PASS" if r["required_pass"] and r["banned_clean"] else "FAIL"
        if status == "FAIL":
            overall_editorial_ok = False
        lines.append(f"### {key} — **{status}**")
        if r["required_misses"]:
            lines.append("- Missing required:")
            for m in r["required_misses"]:
                lines.append(f"  - `{m}`")
        if r["banned_hits"]:
            lines.append("- Banned still present:")
            for m in r["banned_hits"]:
                lines.append(f"  - `{m}`")
        if status == "PASS":
            lines.append(f"- All {len(req)} required present; all {len(ban)} banned absent.")
        lines.append("")

    # --- Section 2: Italics ---
    lines.append("## 2. Italic preservation")
    lines.append("")
    lines.append("### DOCX (BUILD vs ITALICIZED source)")
    lines.append("")
    lines.append("| Book | Source italic runs | BUILD italic runs | Source italic chars | BUILD italic chars | Ratio chars |")
    lines.append("|---|---:|---:|---:|---:|---:|")
    italic_ok = True
    for b in (1, 2, 3):
        src = load_docx_italic_stats(FORMATS[f"SRC_DOCX_{b}"])
        bld = load_docx_italic_stats(FORMATS[f"BUILD_DOCX_{b}"])
        ratio = (bld["italic_chars"] / src["italic_chars"]) if src["italic_chars"] else 0
        # Patch method should retain most italics; flag if < 85%
        flag = "" if ratio >= 0.85 else " ⚠️"
        if ratio < 0.85:
            italic_ok = False
        lines.append(
            f"| {b} | {src['italic_runs']} | {bld['italic_runs']} | "
            f"{src['italic_chars']} | {bld['italic_chars']} | {ratio:.2%}{flag} |"
        )
    lines.append("")
    lines.append("### PDF italic spans (sampled)")
    lines.append("")
    lines.append("| Format | Pages sampled | Italic spans | Italic chars |")
    lines.append("|---|---:|---:|---:|")
    pdf_italic = {}
    for key in ["OMNI_HC", "OMNI_PB", "B1_HC", "B1_PB", "B2_HC", "B2_PB", "B3_HC", "B3_PB"]:
        path = FORMATS[key]
        if path.exists():
            st = load_pdf_italic_stats(path)
            pdf_italic[key] = st
            lines.append(
                f"| {key} | {st['pages_sampled']} | {st['italic_spans']} | {st['italic_chars']} |"
            )
            if st["italic_spans"] == 0:
                italic_ok = False
                lines.append(f"  - ⚠️ No italic spans detected in sample for {key}")
    lines.append("")
    lines.append("### EPUB italic tags")
    lines.append("")
    lines.append("| Format | `<em>` | `<i>` | Total |")
    lines.append("|---|---:|---:|---:|")
    for key in ["B1_EPUB", "B2_EPUB", "B3_EPUB"]:
        path = FORMATS[key]
        if path.exists():
            st = load_epub_italic_stats(path)
            lines.append(f"| {key} | {st['em_tags']} | {st['i_tags']} | {st['italic_tags']} |")
            if st["italic_tags"] == 0:
                italic_ok = False
    lines.append("")

    # --- Section 3: Cross-format consistency ---
    lines.append("## 3. Cross-format consistency")
    lines.append("")
    lines.append("For each book, required strings must appear in BUILD + HC + PB + EPUB (+ Omnibus).")
    lines.append("")
    consistency_ok = True
    for book in (1, 2, 3):
        keys = [f"BUILD_DOCX_{book}", f"B{book}_HC", f"B{book}_PB", f"B{book}_EPUB", "OMNI_HC", "OMNI_PB"]
        lines.append(f"### Book {book}")
        lines.append("")
        lines.append("| Needle | " + " | ".join(keys) + " |")
        lines.append("|---|" + "|".join(["---"] * len(keys)) + "|")
        for needle in REQUIRED[book]:
            row = [needle[:40]]
            for k in keys:
                if k not in texts:
                    row.append("—")
                    consistency_ok = False
                    continue
                ok = norm(needle) in texts[k]
                row.append("Y" if ok else "**N**")
                if not ok:
                    consistency_ok = False
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    # --- Section 4: Typography smoke ---
    lines.append("## 4. Typography / trim smoke")
    lines.append("")
    lines.append("| Format | Pages | Trim (in) | ISBN digits on early pages |")
    lines.append("|---|---:|---|---|")
    isbn_map = {
        "OMNI_HC": "9798295884412",
        "OMNI_PB": "9798256072704",
        "B1_HC": "9798295800801",
        "B1_PB": "9798256008048",
        "B2_HC": "9798295812675",
        "B2_PB": "9798256009953",
        "B3_HC": "9798295812705",
        "B3_PB": "9798256010072",
    }
    for key, isbn in isbn_map.items():
        path = FORMATS[key]
        if not path.exists():
            lines.append(f"| {key} | MISSING | | |")
            continue
        doc = fitz.open(path)
        r = doc[0].rect
        front = "".join(doc[i].get_text() for i in range(min(8, doc.page_count)))
        digits = re.sub(r"\D", "", front)
        isbn_ok = isbn in digits
        lines.append(
            f"| {key} | {doc.page_count} | {r.width/72:.2f}×{r.height/72:.2f} | "
            f"{'Y' if isbn_ok else '**N**'} |"
        )
        doc.close()
    lines.append("")

    # Verdict
    lines.append("## Verdict")
    lines.append("")
    if overall_editorial_ok and italic_ok and consistency_ok:
        lines.append("**UPLOAD AUDIT: PASS** — editorial fixes present, banned strings absent, italics detected, formats consistent.")
        verdict = 0
    else:
        lines.append("**UPLOAD AUDIT: FAIL / NEEDS REVIEW**")
        if not overall_editorial_ok:
            lines.append("- Editorial required/banned checks failed on one or more delivery formats.")
        if not italic_ok:
            lines.append("- Italic preservation below threshold or missing in a delivery format.")
        if not consistency_ok:
            lines.append("- Cross-format consistency gaps (see matrix).")
        verdict = 1
    lines.append("")
    lines.append("*Seventh City Press · PRE-UPLOAD AUDIT · f = 111.2 Hz*")

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {OUT}")
    print("editorial_ok", overall_editorial_ok, "italic_ok", italic_ok, "consistency_ok", consistency_ok)
    return verdict


if __name__ == "__main__":
    raise SystemExit(main())
