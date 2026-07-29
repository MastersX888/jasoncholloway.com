#!/usr/bin/env python3
"""
COMPOSITOR Phase A — patch ITALICIZED DOCX -> BUILD docx.

Applies geo + apartment/balcony string replacements while preserving
run formatting when the old string sits inside a single run; otherwise
rewrites the paragraph using the first run's style.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
SOURCES = ROOT / "production_staging" / "_sources"
OUT_DIR = SOURCES / "build_docx"

# (old, new) — short unique fragments preferred; longer blocks when needed.
# Order: longer / more specific first within each book.

REPLACEMENTS: dict[int, list[tuple[str, str]]] = {
    1: [
        (
            '"Picture Cave," she said. "Washington County, Missouri. Dated to '
            'at least a thousand years ago. Pre-Columbian. Pre-European. '
            'Pre-Christian."',
            '"Picture Cave," she said. "Warren County, Missouri. Osage ancestral '
            'ground, sacred to them still. Dated to at least a thousand years ago. '
            'Pre-Columbian. Pre-European. Pre-Christian."',
        ),
        (
            "2847 Genessee Street looked abandoned. Graffiti on the loading "
            "dock. Windows covered with plywood. A faded sign: MISSOURI "
            "COLD STORAGE CO. 1923.",
            "1647 Genessee Street looked abandoned. Graffiti on the loading "
            "dock. Windows covered with plywood. A faded sign: RIVERWARDS "
            "COLD STORAGE CO. 1923.",
        ),
        (
            '"SubTropolis wasn\'t built. It was discovered. The Bethany Falls '
            'Limestone Company found existing cave systems in the 1940s. The '
            'caves Moreau used were there decades before. Centuries before."',
            '"SubTropolis wasn\'t built. It was mined. Midwest Precote started '
            'room-and-pillar work in the Bethany Falls limestone in the 1940s, and '
            'Hunt Midwest turned the emptied rooms into an underground business '
            'district twenty years later. But the mining broke into voids nobody '
            'cut. The caves Moreau used were there decades before. Centuries '
            'before."',
        ),
        (
            "The Analysis Chamber lived in the basement of 2847 Genessee Street.",
            "The Analysis Chamber lived in the basement of 1647 Genessee Street.",
        ),
        (
            "The commercial mining level, SubTropolis proper, fifty-five feet "
            "below the surface, ended here. Everything below this point was "
            "older, deeper, and no longer part of anyone's inventory. The shaft "
            "dropped a hundred and sixty feet into limestone that had been "
            "forming since before anything human existed.",
            "The commercial mining level, SubTropolis proper, a hundred and "
            "fifty feet below the surface, ended here. Everything below this "
            "point was older, deeper, and no longer part of anyone's inventory. "
            "The shaft bottomed out at a hundred and sixty feet, in limestone "
            "that had been forming since before anything human existed.",
        ),
    ],
    2: [
        (
            "buildings of Troost Avenue with the indifferent clarity",
            "buildings of Pennsylvania Avenue with the indifferent clarity",
        ),
        (
            "outside the Quality Hill entrance on graduation day, twenty-three faces,",
            "outside the Washington Street entrance on graduation day, twenty-three faces,",
        ),
        (
            "She walked to the Foundation's mailbox. Quality Hill. February.",
            "She walked to the Foundation's mailbox. Washington Street. February.",
        ),
        (
            "The mailbox on Quality Hill held a letter",
            "The mailbox on Washington Street held a letter",
        ),
        (
            "11 PM. Quality Hill. The Foundation empty.",
            "11 PM. Washington Street. The Foundation empty.",
        ),
        (
            "The Hotel Phillips Building office in November.",
            "The Washington Street office in November.",
        ),
        (
            "the stone cottage and the Hotel Phillips Building office,",
            "the stone cottage and the Washington Street office,",
        ),
        (
            "The house hummed its limestone fundamental,",
            "The apartment hummed its limestone fundamental,",
        ),
        (
            "bathroom sink in a house in Kansas City.",
            "bathroom sink in an apartment in Kansas City.",
        ),
    ],
    3: [
        # Geo
        (
            "333.6 Hz harmonic in the basalt chamber. She was Cohort 4",
            "333.6 Hz harmonic in the Iceland basalt chamber. She was Cohort 4",
        ),
        (
            "The basalt chamber resonance is stable at 111.2.",
            "The Iceland basalt chamber is stable at 111.2.",
        ),
        (
            "Somewhere below them, 160 feet below, in the SubTropolis",
            "Across town, 160 feet below the Northland bluff, in the SubTropolis",
        ),
        (
            "the streets of the Troost corridor empty, the industrial district's daytime commerce",
            "the streets of Quality Hill empty, the district's daytime commerce",
        ),
        (
            "the industrial memory of the West Bottoms, where, 160 feet below "
            "ground, the Moreau chamber sat in its limestone silence",
            "the industrial memory of the West Bottoms below the bluff. Seven miles "
            "northeast, across the river at SubTropolis, 160 feet down, the Moreau "
            "chamber sat in its limestone silence",
        ),
        (
            "She drove down Troost. Past the gas station. Past the church "
            "where they'd held the first community protocol session. Past the "
            "Quality Hill boundary into the institutional district.",
            "She drove the long way. East to Troost. Past the gas station. Past the church "
            "where they'd held the first community protocol session. Then west on "
            "Twelfth, across Broadway, up the bluff to Washington Street and the "
            "institutional district.",
        ),
        (
            "second floor of the Washington Street building, the room where Nadia",
            "third floor of the Washington Street building, the room where Nadia",
        ),
        (
            "West Bottoms, the underground chambers beneath the bluffs, and felt",
            "West Bottoms, and across the river the underground chambers at SubTropolis, and felt",
        ),
        (
            "Marcus Chen on his front porch, wearing a Kansas City Royals cap",
            "Marcus Chen in the hallway outside his door, wearing a Kansas City Royals cap",
        ),
        # Apartment / balcony (high-signal unique phrases)
        ("The house was clean. The house was empty.",
         "The apartment was clean. The apartment was empty."),
        ("The house had a frequency.", "The apartment had a frequency."),
        ("this particular house under this particular life.",
         "this particular apartment under this particular life."),
        ("sink in this house, on this counter.",
         "sink in this apartment, on this counter."),
        ("Just a man, in a house, at a table",
         "Just a man, in an apartment, at a table"),
        ("the kitchen and the house and the",
         "the kitchen and the apartment and the"),
        ("foundation beneath the house and the Kansas City limestone",
         "foundation beneath the apartment and the Kansas City limestone"),
        ('"I can hear the house."', '"I can hear the apartment."'),
        ("The house has a frequency, the composite",
         "The apartment has a frequency, the composite"),
        ("The house sounds like an old instrument.",
         "The apartment sounds like an old instrument."),
        ('"I don\'t know if I\'m talking about the house."',
         '"I don\'t know if I\'m talking about the apartment."'),
        ("The house was quiet. The house was tired. The house kept",
         "The apartment was quiet. The apartment was tired. The apartment kept"),
        ("of your house, not by measuring", "of your home, not by measuring"),
        ("the house humming its 55 Hz", "the apartment humming its 55 Hz"),
        ("to anyone in this house, you hear",
         "to anyone in this apartment, you hear"),
        ("limestone under this house.", "limestone under this apartment."),
        ("Not in this house.", "Not in this apartment."),
        ("leaving this house at", "leaving this apartment at"),
        ("The house hummed at 55 Hz.", "The apartment hummed at 55 Hz."),
        ("sound of the house.", "sound of the apartment."),
        ("contribution to the sound of the house.",
         "contribution to the sound of the apartment."),
        ("her system was the house and the",
         "her system was the apartment and the"),
        ("house was still running on her architecture",
         "apartment was still running on her architecture"),
        ("and the house would recalibrate",
         "and the apartment would recalibrate"),
        ("The house hummed at 53.6 Hz, its new pitch, the sound of a house with one heartbeat",
         "The apartment hummed at 53.6 Hz, its new pitch, the sound of an apartment with one heartbeat"),
        ("He sat on the Quality Hill porch.", "He sat on the Quality Hill balcony."),
        ("The porch was old, the original porch, the porch his grandfather",
         "The balcony was old, the original balcony, the balcony his grandfather"),
        ("when they lean on porches and look",
         "when they lean on balconies and look"),
        ("first night in this house", "first night in this apartment"),
        ("He fixed the porch railing.", "He fixed the balcony railing."),
        ("into the wood into the house. He",
         "into the wood into the building. He"),
        ("was shaking the house. Very slightly.",
         "was shaking the building. Very slightly."),
        ("the house, unless you had limestone foundations",
         "the building, unless you had limestone foundations"),
        ("fixing a porch rail.", "fixing a balcony rail."),
        ("She found him on the porch.", "She found him on the balcony."),
        ("on the porch rail, the rail he", "on the balcony rail, the rail he"),
        ("sitting on a porch drinking wine",
         "sitting on a balcony drinking wine"),
        ("Day fifteen. The porch.", "Day fifteen. The balcony."),
        ("on a porch in October", "on a balcony in October"),
        ("frequency of the house's own limestone foundation",
         "frequency of the building's own limestone foundation"),
        ("this particular porch drinking this particular wine.",
         "this particular balcony drinking this particular wine."),
        ("the light moved across the porch",
         "the light moved across the balcony"),
        ("porch at 9 PM.", "balcony at 9 PM."),
        ("entering a house that had been her house and was her house again.",
         "entering an apartment that had been her apartment and was her apartment again."),
        ("separate bedrooms and separate houses and",
         "separate bedrooms and separate apartments and"),
        ('And listen to the house."', 'And listen to the apartment."'),
        # Catch remaining "The house's" possessive if still present
        ("The house's", "The apartment's"),
        ("geological signature of this particular plot of earth under this particular house under",
         "geological signature of this particular plot of earth under this particular apartment under"),
    ],
}

VERIFY = {
    1: ["1647 Genessee", "Warren County", "Midwest Precote", "hundred and fifty feet"],
    2: ["Washington Street office", "Pennsylvania Avenue", "apartment hummed"],
    3: [
        "across the river at SubTropolis",
        "Iceland basalt",
        "Quality Hill balcony",
        "apartment hummed",
        "the long way",
    ],
}


def plain(p: Paragraph) -> str:
    if p.runs:
        return "".join(r.text or "" for r in p.runs)
    return p.text or ""


def rewrite_para(p: Paragraph, new_text: str) -> None:
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def replace_once_in_doc(doc: Document, old: str, new: str) -> str:
    """Return 'hit', 'already', or 'miss'."""
    # Fast scan with p.text (cached by python-docx)
    for p in doc.paragraphs:
        t = p.text or ""
        if old in t:
            # Prefer single-run replace
            for r in p.runs:
                if old in (r.text or ""):
                    r.text = (r.text or "").replace(old, new, 1)
                    return "hit"
            rewrite_para(p, t.replace(old, new, 1))
            return "hit"
        if new in t and old not in t:
            # might already be applied elsewhere; keep scanning
            continue
    # Already applied?
    full = "\n".join(p.text or "" for p in doc.paragraphs)
    if new in full and old not in full:
        return "already"
    # Try curly-apostrophe / curly-quote variants of old
    variants = [
        old.replace("'", "\u2019"),
        old.replace("'", "\u2019").replace('"', "\u201c", 1),
        old.replace("'", "\u2019").replace('"', "\u201c").replace('"', "\u201d"),
    ]
    for var in variants:
        if var == old:
            continue
        for p in doc.paragraphs:
            t = p.text or ""
            if var in t:
                for r in p.runs:
                    if var in (r.text or ""):
                        r.text = (r.text or "").replace(var, new, 1)
                        return "hit"
                rewrite_para(p, t.replace(var, new, 1))
                return "hit"
    return "miss"


def build_book(book: int) -> tuple[Path, bool]:
    src = SOURCES / f"MASTERS_X_BOOK{book}_ITALICIZED_FIXED.docx"
    print(f"Loading Book {book}: {src.name}", flush=True)
    doc = Document(str(src))
    hits = already = misses = 0
    miss_list = []
    for old, new in REPLACEMENTS[book]:
        status = replace_once_in_doc(doc, old, new)
        if status == "hit":
            hits += 1
        elif status == "already":
            already += 1
        else:
            misses += 1
            miss_list.append(old[:70])

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"MASTERS_X_BOOK{book}_BUILD.docx"
    doc.save(str(out))
    print(f"  saved {out.name}  hits={hits} already={already} misses={misses}", flush=True)
    for m in miss_list:
        print(f"    MISS: {m!r}", flush=True)

    full = "\n".join(p.text or "" for p in doc.paragraphs)
    ok = True
    print("  VERIFY:", flush=True)
    for needle in VERIFY[book]:
        present = needle in full
        print(f"    {needle!r} -> {present}", flush=True)
        ok = ok and present
    # Extra Book 3 checks from handover
    if book == 3:
        for needle in ["She drove", "balcony"]:
            print(f"    extra {needle!r} -> {needle in full}", flush=True)
    return out, ok


def main() -> int:
    print("COMPOSITOR Phase A — BUILD docx", flush=True)
    results = []
    for b in (1, 2, 3):
        path, ok = build_book(b)
        results.append((b, path, ok))
    print("\nSUMMARY", flush=True)
    for b, path, ok in results:
        print(f"  Book {b}: {'PASS' if ok else 'FAIL'}  {path}", flush=True)
    if not all(r[2] for r in results):
        print("STOP: BUILD verify failed.", flush=True)
        return 1
    print("All BUILD docx verified.", flush=True)
    return 0


if __name__ == "__main__":
    sys.exit(main())
